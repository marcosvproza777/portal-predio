"""
Gerador de Relatório Executivo Word — Pred.IO (Etapa 6.5)
Gera .docx editável com histórico completo de um ativo.
"""
from __future__ import annotations
import io
import datetime

# ── Paleta Pred.IO (hex sem #) ────────────────────────────────────────────────
_NAVY   = "0F1F3D"
_BLUE   = "2563EB"
_GREEN  = "10B981"
_AMBER  = "F59E0B"
_RED    = "EF4444"
_MUTED  = "64748B"
_LGRAY  = "F1F5F9"
_WHITE  = "FFFFFF"
_BORDER = "CBD5E1"

# ── Tuple RGB helpers ─────────────────────────────────────────────────────────
def _rgb(h: str):
    from docx.shared import RGBColor
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


# ═══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════
def gerar_relatorio_word(dados: dict) -> bytes:
    """Gera o relatório executivo Word e retorna bytes do .docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx não instalado. Verifique requirements.txt.")

    doc = Document()
    _cfg_pagina(doc)

    _secao_capa(doc, dados)
    doc.add_page_break()
    _secao_resumo_executivo(doc, dados)
    doc.add_page_break()
    _secao_identificacao(doc, dados)
    _secao_indicadores(doc, dados)

    if dados.get("incluir_graficos", True):
        doc.add_page_break()
        _secao_graficos(doc, dados)

    if dados.get("incluir_historico", True):
        doc.add_page_break()
        _secao_historico(doc, dados)

    if dados.get("incluir_manutencoes", True):
        doc.add_page_break()
        _secao_manutencoes_realizadas(doc, dados)

    if dados.get("incluir_manutencoes_futuras", True):
        doc.add_page_break()
        _secao_manutencoes_pendentes(doc, dados)

    if dados.get("incluir_relatorios", True):
        doc.add_page_break()
        _secao_relatorios(doc, dados)

    doc.add_page_break()
    _secao_principais_pontos(doc, dados)

    if dados.get("incluir_alertas", True):
        doc.add_page_break()
        _secao_alertas_chamados(doc, dados)

    doc.add_page_break()
    _secao_recomendacoes(doc, dados)

    if dados.get("incluir_recomendacoes_condicao", True):
        doc.add_page_break()
        _secao_recomendacoes_condicao(doc, dados)

    doc.add_page_break()
    _secao_conclusao(doc, dados)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURAÇÃO DE PÁGINA
# ═══════════════════════════════════════════════════════════════════════════════
def _cfg_pagina(doc) -> None:
    from docx.shared import Cm
    s = doc.sections[0]
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(2.8)
    s.right_margin  = Cm(2.8)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS VISUAIS
# ═══════════════════════════════════════════════════════════════════════════════
def _hd(doc, texto: str, nivel: int = 1) -> None:
    """Cabeçalho de seção numerada."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if nivel == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(texto)
    run.bold = True
    run.font.size  = Pt(13 if nivel == 1 else 11)
    run.font.color.rgb = _rgb(_NAVY if nivel == 1 else _BLUE)
    if nivel == 1:
        p.paragraph_format.left_indent = Pt(0)
        # linha divisória abaixo via borda inferior
        _set_para_bottom_border(p)


def _txt(doc, texto: str, size_pt: int = 10, bold: bool = False,
         color: str = "2D3748", italic: bool = False) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto)
    run.bold   = bold
    run.italic = italic
    run.font.size      = Pt(size_pt)
    run.font.color.rgb = _rgb(color)


def _tbl(doc, headers: list[str], rows: list[list], col_widths_cm: list[float] | None = None):
    """Cria tabela com cabeçalho azul escuro e linhas alternadas."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"

    # cabeçalho
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, _NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size      = Pt(8.5)
        run.font.color.rgb = _rgb(_WHITE)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # dados
    for ri, row in enumerate(rows):
        bg = _WHITE if ri % 2 == 0 else _LGRAY
        trow = table.rows[1 + ri]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "—")
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # larguras
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    _set_tbl_borders(table)
    doc.add_paragraph()
    return table


def _badge_cell(cell, texto: str, cor_hex: str) -> None:
    """Texto colorido numa célula de tabela."""
    from docx.shared import Pt
    _set_cell_bg(cell, cor_hex + "22")
    p = cell.paragraphs[0]
    run = p.add_run(texto)
    run.bold = True
    run.font.size      = Pt(8.5)
    run.font.color.rgb = _rgb(cor_hex)


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_tbl_borders(table) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), _BORDER)
        borders.append(b)
    tblPr.append(borders)


def _set_para_bottom_border(p) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), _NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# GRÁFICOS (matplotlib → PNG → Word)
# ═══════════════════════════════════════════════════════════════════════════════
def _grafico_bytes(fig) -> io.BytesIO:
    import matplotlib.pyplot as plt
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _grafico_score(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np
    except ImportError:
        return None

    ativo = dados.get("ativo", {})
    try:
        score = int(float(str(ativo.get("Score", 0))))
    except (ValueError, TypeError):
        score = 0

    fig, ax = plt.subplots(figsize=(5, 2.5), facecolor="white")
    cor = "#10B981" if score >= 85 else "#F59E0B" if score >= 60 else "#EF4444"
    ax.barh(["Score de Saúde"], [score], color=cor, height=0.45, zorder=3)
    ax.barh(["Score de Saúde"], [100 - score], left=[score], color="#E2E8F0", height=0.45, zorder=2)
    ax.set_xlim(0, 100)
    ax.set_xlabel("Pontos (0–100)", fontsize=9, color="#475569")
    ax.set_title("Score de Saúde Atual", fontsize=11, fontweight="bold", color="#0F1F3D", pad=8)
    ax.text(score + 1 if score < 90 else score - 8, 0, f"{score}/100",
            va="center", fontsize=11, fontweight="bold", color=cor)
    ax.tick_params(left=False, labelleft=False, colors="#475569", labelsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.grid(axis="x", linestyle="--", alpha=0.4, color="#CBD5E1", zorder=1)
    fig.tight_layout()
    return _grafico_bytes(fig)


def _grafico_relatorios_severidade(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    df = dados.get("relatorios")
    if df is None or df.empty:
        return None

    cats  = ["Normal", "Atenção", "Crítico", "Urgente"]
    cores = ["#10B981", "#F59E0B", "#EF4444", "#7C3AED"]
    col   = next((c for c in ("Severidade", "Status", "Tipo_Servico") if c in df.columns), None)
    counts = []
    for cat in cats:
        if col:
            counts.append(int(df[col].str.lower().str.contains(cat.lower(), na=False).sum()))
        else:
            counts.append(0)

    if sum(counts) == 0:
        return None

    fig, ax = plt.subplots(figsize=(4.5, 3), facecolor="white")
    wedges, texts, autotexts = ax.pie(
        counts, labels=cats, colors=cores, autopct=lambda p: f"{p:.0f}%" if p > 0 else "",
        startangle=90, pctdistance=0.75,
    )
    for t in texts:
        t.set_fontsize(9)
        t.set_color("#475569")
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color("white")
        at.set_fontweight("bold")
    ax.set_title("Relatórios por Severidade", fontsize=11, fontweight="bold",
                 color="#0F1F3D", pad=8)
    fig.tight_layout()
    return _grafico_bytes(fig)


def _grafico_manutencoes(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
        import numpy as np
    except ImportError:
        return None

    df_pend = dados.get("manutencoes_pendentes")
    if df_pend is None or df_pend.empty:
        return None

    cats  = ["Em dia", "Próx. vencimento", "Vencidas", "Por condição"]
    cores = ["#10B981", "#F59E0B", "#EF4444", "#6366F1"]
    counts = [0, 0, 0, 0]

    if "Status_Calculado" in df_pend.columns:
        for i, row in df_pend.iterrows():
            s = str(row.get("Status_Calculado", "")).lower()
            if "vencid" in s or "atraso" in s:
                counts[2] += 1
            elif "próxim" in s or "proxi" in s or "próximo" in s:
                counts[1] += 1
            elif "condi" in s:
                counts[3] += 1
            else:
                counts[0] += 1
    elif "Tipo_Manutencao" in df_pend.columns:
        for i, row in df_pend.iterrows():
            tipo = str(row.get("Tipo_Manutencao", "")).lower()
            if "condi" in tipo:
                counts[3] += 1
            else:
                counts[0] += 1

    if sum(counts) == 0:
        return None

    fig, ax = plt.subplots(figsize=(5, 2.5), facecolor="white")
    bars = ax.bar(cats, counts, color=cores, width=0.55, zorder=3)
    ax.set_title("Manutenções por Status", fontsize=11, fontweight="bold",
                 color="#0F1F3D", pad=8)
    ax.set_ylabel("Quantidade", fontsize=9, color="#475569")
    ax.tick_params(colors="#475569", labelsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4, color="#CBD5E1", zorder=1)
    for bar in bars:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.1, str(int(h)),
                    ha="center", va="bottom", fontsize=9, fontweight="bold", color="#334155")
    fig.tight_layout()
    return _grafico_bytes(fig)


def _grafico_chamados(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    df = dados.get("chamados")
    if df is None or df.empty:
        return None

    cats  = ["Aberto", "Em análise", "Aguardando", "Concluído"]
    cores = ["#3B82F6", "#F59E0B", "#6366F1", "#10B981"]
    col   = next((c for c in ("Status",) if c in df.columns), None)
    counts = []
    for cat in cats:
        if col:
            counts.append(int(df[col].str.lower().str.contains(cat.lower().replace("í","i"), na=False).sum()))
        else:
            counts.append(0)

    if sum(counts) == 0:
        return None

    fig, ax = plt.subplots(figsize=(5, 2.5), facecolor="white")
    bars = ax.bar(cats, counts, color=cores, width=0.55, zorder=3)
    ax.set_title("Chamados por Status", fontsize=11, fontweight="bold",
                 color="#0F1F3D", pad=8)
    ax.set_ylabel("Quantidade", fontsize=9, color="#475569")
    ax.tick_params(colors="#475569", labelsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4, color="#CBD5E1", zorder=1)
    for bar in bars:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.05, str(int(h)),
                    ha="center", va="bottom", fontsize=9, fontweight="bold", color="#334155")
    fig.tight_layout()
    return _grafico_bytes(fig)


def _grafico_alertas(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    df = dados.get("alertas")
    if df is None or df.empty:
        return None

    cats  = ["Baixa", "Média", "Alta", "Crítica"]
    cores = ["#10B981", "#F59E0B", "#EF4444", "#7C3AED"]
    col   = next((c for c in ("Prioridade", "Severidade") if c in df.columns), None)
    counts = []
    for cat in cats:
        if col:
            counts.append(int(df[col].str.lower().str.contains(cat.lower(), na=False).sum()))
        else:
            counts.append(0)

    if sum(counts) == 0:
        return None

    fig, ax = plt.subplots(figsize=(5, 2.5), facecolor="white")
    bars = ax.bar(cats, counts, color=cores, width=0.55, zorder=3)
    ax.set_title("Alertas por Prioridade", fontsize=11, fontweight="bold",
                 color="#0F1F3D", pad=8)
    ax.set_ylabel("Quantidade", fontsize=9, color="#475569")
    ax.tick_params(colors="#475569", labelsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4, color="#CBD5E1", zorder=1)
    for bar in bars:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.05, str(int(h)),
                    ha="center", va="bottom", fontsize=9, fontweight="bold", color="#334155")
    fig.tight_layout()
    return _grafico_bytes(fig)


def _grafico_eventos_mes(dados: dict) -> io.BytesIO | None:
    try:
        import matplotlib.pyplot as plt
        import numpy as np
    except ImportError:
        return None

    import pandas as pd

    meses_dados: dict[str, dict[str, int]] = {}

    def _add(df, col_data: str, label: str):
        if df is None or df.empty or col_data not in df.columns:
            return
        for val in df[col_data].dropna():
            mes = str(val)[:7]
            if len(mes) < 7:
                continue
            meses_dados.setdefault(mes, {"Relatórios": 0, "Manutenções": 0, "Alertas": 0, "Chamados": 0})
            meses_dados[mes][label] += 1

    _add(dados.get("relatorios"),             "Data_Relatorio",  "Relatórios")
    _add(dados.get("manutencoes_executadas"),  "Executado_Em",    "Manutenções")
    _add(dados.get("alertas"),                 "Data",            "Alertas")
    _add(dados.get("chamados"),                "Data_Abertura",   "Chamados")

    if not meses_dados:
        return None

    meses = sorted(meses_dados.keys())[-12:]
    cats  = ["Relatórios", "Manutenções", "Alertas", "Chamados"]
    cores = ["#2563EB", "#10B981", "#EF4444", "#F59E0B"]

    x     = np.arange(len(meses))
    width = 0.18
    fig, ax = plt.subplots(figsize=(7, 3), facecolor="white")
    for i, (cat, cor) in enumerate(zip(cats, cores)):
        vals = [meses_dados.get(m, {}).get(cat, 0) for m in meses]
        ax.bar(x + i * width - 1.5 * width, vals, width, label=cat, color=cor, zorder=3)

    ax.set_title("Eventos Técnicos por Mês", fontsize=11, fontweight="bold",
                 color="#0F1F3D", pad=8)
    ax.set_xticks(x)
    ax.set_xticklabels([m[5:] + "/" + m[2:4] for m in meses], fontsize=8, rotation=30, color="#475569")
    ax.set_ylabel("Quantidade", fontsize=9, color="#475569")
    ax.tick_params(colors="#475569", labelsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4, color="#CBD5E1", zorder=1)
    ax.legend(fontsize=8, framealpha=0.8)
    fig.tight_layout()
    return _grafico_bytes(fig)


def _inserir_grafico(doc, buf: io.BytesIO | None, legenda: str = "") -> None:
    if buf is None:
        return
    from docx.shared import Inches
    doc.add_picture(buf, width=Inches(6.0))
    if legenda:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = doc.add_paragraph(legenda)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.italic   = True
        run.font.size      = Pt(8)
        run.font.color.rgb = _rgb(_MUTED)


# ═══════════════════════════════════════════════════════════════════════════════
# SEÇÕES DO RELATÓRIO
# ═══════════════════════════════════════════════════════════════════════════════
def _secao_capa(doc, dados: dict) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ativo        = dados.get("ativo", {})
    nome_ativo   = str(ativo.get("Tag", "")).strip() or str(ativo.get("Nome", "")).strip() or "Ativo"
    cliente_nome = dados.get("cliente_nome", "")
    periodo_ini  = dados.get("periodo_inicio", "")
    periodo_fim  = dados.get("periodo_fim", "")
    data_geracao = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    # espaçamento superior
    for _ in range(6):
        doc.add_paragraph()

    def _capa_par(texto, tamanho, negrito=False, cor=_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(texto)
        run.bold           = negrito
        run.font.size      = Pt(tamanho)
        run.font.color.rgb = _rgb(cor)
        return p

    _capa_par("Pred.IO", 28, negrito=True, cor=_NAVY)
    _capa_par("Relatório Executivo de Confiabilidade", 18, negrito=True, cor=_BLUE)
    doc.add_paragraph()
    _capa_par(nome_ativo, 16, negrito=True, cor=_NAVY)
    _capa_par(f"Cliente: {cliente_nome}", 12, cor=_MUTED)

    doc.add_paragraph()

    if periodo_ini and periodo_fim:
        _capa_par(f"Período analisado: {periodo_ini} a {periodo_fim}", 11, cor=_MUTED)
    _capa_par(f"Data de geração: {data_geracao}", 10, cor=_MUTED)
    _capa_par("Fonte: Pred.IO", 10, cor=_MUTED)

    doc.add_paragraph()
    _capa_par(
        "Documento editável — Para uso interno Pred.IO. "
        "Revisar antes do envio ao cliente.",
        9, italic_=True if False else False, cor="DC2626",
    )


def _secao_resumo_executivo(doc, dados: dict) -> None:
    _hd(doc, "1. Sumário Executivo")
    texto = _gerar_texto_resumo(dados)
    _txt(doc, texto, size_pt=10)

    # Principais riscos e recomendações como lista
    _hd(doc, "Principais pontos:", nivel=2)
    pontos = _gerar_pontos_resumo(dados)
    for pt in pontos:
        p = doc.add_paragraph(style="List Bullet")
        from docx.shared import Pt
        run = p.add_run(pt)
        run.font.size = Pt(9.5)


def _secao_identificacao(doc, dados: dict) -> None:
    _hd(doc, "2. Identificação do Ativo")
    ativo = dados.get("ativo", {})
    horimetro = dados.get("horimetro")

    rows = [
        ["Nome do ativo",      str(ativo.get("Tag", "")).strip() or str(ativo.get("Nome", "")).strip()],
        ["Cliente",            dados.get("cliente_nome", "")],
        ["Planta / Setor",     str(ativo.get("Planta", "")).strip()],
        ["Tipo de equipamento",str(ativo.get("Tipo", "")).strip()],
        ["Fabricante / Modelo",str(ativo.get("Modelo", "")).strip()],
        ["Número de série",    str(ativo.get("Ns", "") or ativo.get("Numero_Serie", "")).strip()],
        ["MB",                 str(ativo.get("Mb", "")).strip()],
        ["Inversor de freq.",  str(ativo.get("Inversor", "")).strip()],
        ["Análise de óleo",    str(ativo.get("Analise_Oleo", "")).strip()],
        ["Horímetro atual",    f"{horimetro} h" if horimetro else "Não informado"],
        ["Criticidade",        str(ativo.get("Criticidade", "")).strip()],
        ["Status de saúde",    str(ativo.get("Status", "")).strip()],
        ["Score de saúde",     str(ativo.get("Score", "")).strip() + "/100"],
        ["Última atualização", str(ativo.get("Data", "")).strip()],
    ]
    _tbl(doc, ["Campo", "Valor"], rows, col_widths_cm=[5.5, 10.5])


def _secao_indicadores(doc, dados: dict) -> None:
    _hd(doc, "3. Indicadores Principais do Período")
    df_cham  = dados.get("chamados")
    df_al    = dados.get("alertas")
    df_rel   = dados.get("relatorios")
    df_mex   = dados.get("manutencoes_executadas")
    df_mpend = dados.get("manutencoes_pendentes")

    def _cnt(df): return 0 if df is None or df.empty else len(df)
    def _cnt_col(df, col, val):
        if df is None or df.empty or col not in df.columns:
            return 0
        return int(df[col].str.lower().str.contains(val.lower(), na=False).sum())

    n_rel    = _cnt(df_rel)
    n_mex    = _cnt(df_mex)
    n_mpend  = _cnt(df_mpend)
    n_mvenc  = _cnt_col(df_mpend, "Status_Calculado", "vencid") if df_mpend is not None else 0
    n_cham   = _cnt(df_cham)
    n_al     = _cnt(df_al)
    n_al_crit= _cnt_col(df_al, "Prioridade", "crítica") if df_al is not None else 0

    rows = [
        ["Relatórios técnicos no período", str(n_rel)],
        ["Manutenções executadas",          str(n_mex)],
        ["Manutenções pendentes (total)",   str(n_mpend)],
        ["Manutenções vencidas",            str(n_mvenc)],
        ["Alertas registrados",             str(n_al)],
        ["Alertas críticos",                str(n_al_crit)],
        ["Chamados técnicos",               str(n_cham)],
    ]
    _tbl(doc, ["Indicador", "Quantidade"], rows, col_widths_cm=[11.0, 5.0])


def _secao_graficos(doc, dados: dict) -> None:
    _hd(doc, "4. Gráficos Executivos")

    graficos = [
        (_grafico_score(dados),                    "Gráfico 1 — Score de Saúde Atual"),
        (_grafico_relatorios_severidade(dados),    "Gráfico 2 — Relatórios por Severidade"),
        (_grafico_manutencoes(dados),              "Gráfico 3 — Manutenções por Status"),
        (_grafico_chamados(dados),                 "Gráfico 4 — Chamados por Status"),
        (_grafico_alertas(dados),                  "Gráfico 5 — Alertas por Prioridade"),
        (_grafico_eventos_mes(dados),              "Gráfico 6 — Eventos Técnicos por Mês"),
    ]
    gerados = 0
    for buf, legenda in graficos:
        if buf is not None:
            _inserir_grafico(doc, buf, legenda)
            gerados += 1

    if gerados == 0:
        _txt(doc, "Dados insuficientes para gerar gráficos no período selecionado.", cor=_MUTED)


def _secao_historico(doc, dados: dict) -> None:
    import pandas as pd
    _hd(doc, "5. Histórico Técnico do Ativo")
    _txt(doc, "Linha do tempo cronológica de eventos técnicos registrados no período.", cor=_MUTED, size_pt=9)

    eventos: list[list] = []

    def _add_ev(df, col_data, col_titulo, tipo, col_sev=None):
        if df is None or df.empty:
            return
        for _, row in df.iterrows():
            data  = str(row.get(col_data, "")).strip()[:10]
            tit   = str(row.get(col_titulo, "")).strip()
            sev   = str(row.get(col_sev, "")).strip() if col_sev else "—"
            eventos.append([data, tipo, tit[:80], sev])

    _add_ev(dados.get("relatorios"),            "Data_Relatorio",  "Titulo",   "Relatório Técnico",   "Severidade")
    _add_ev(dados.get("manutencoes_executadas"), "Executado_Em",    "Tarefa",   "Manutenção Executada","Tipo_Manutencao")
    _add_ev(dados.get("chamados"),               "Data_Abertura",   "Titulo",   "Chamado Técnico",     "Prioridade")
    _add_ev(dados.get("alertas"),                "Data",            "Titulo",   "Alerta",              "Prioridade")

    eventos.sort(key=lambda x: x[0] or "", reverse=True)

    if not eventos:
        _txt(doc, "Nenhum evento registrado no período selecionado.", cor=_MUTED)
        return

    _tbl(doc, ["Data", "Tipo de Evento", "Título", "Severidade / Prioridade"],
         eventos, col_widths_cm=[2.5, 4.0, 7.5, 4.0])


def _secao_manutencoes_realizadas(doc, dados: dict) -> None:
    _hd(doc, "6. Manutenções Realizadas")
    df = dados.get("manutencoes_executadas")
    if df is None or df.empty:
        _txt(doc, "Nenhuma manutenção executada registrada no período.", cor=_MUTED)
        return

    rows = []
    for _, row in df.iterrows():
        rows.append([
            str(row.get("Executado_Em",      "")).strip()[:10],
            str(row.get("Tarefa",            "")).strip()[:60],
            str(row.get("Tipo_Manutencao",   "")).strip(),
            str(row.get("Responsavel",       "")).strip(),
            str(row.get("Horimetro_Execucao","")).strip(),
            str(row.get("Obs_Visivel",       "")).strip()[:80],
        ])

    _tbl(doc,
         ["Data", "Tarefa", "Tipo", "Responsável", "Horímetro", "Observação"],
         rows, col_widths_cm=[2.5, 5.0, 3.0, 3.5, 2.5, 5.5])


def _secao_manutencoes_pendentes(doc, dados: dict) -> None:
    _hd(doc, "7. Manutenções a Realizar")
    _txt(doc,
         "⚠️ Overhaul, kit revisão, desmontagem e troca de rolamento aparecem apenas como "
         "'Depende de análise preditiva' — nunca como obrigação automática por horímetro.",
         size_pt=8.5, italic=True, color=_AMBER)
    df = dados.get("manutencoes_pendentes")
    if df is None or df.empty:
        _txt(doc, "Nenhuma tarefa de manutenção pendente encontrada.", cor=_MUTED)
        return

    rows = []
    for _, row in df.iterrows():
        tarefa = str(row.get("Tarefa", "")).strip()
        tipo   = str(row.get("Tipo_Manutencao", "")).strip()
        prox_d = str(row.get("Proxima_Execucao_Data", "")).strip()[:10]
        prox_h = str(row.get("Proxima_Execucao_Horimetro", "")).strip()
        status = str(row.get("Status_Calculado", "Pendente")).strip()

        # Regra de negócio: overhaul/rolamento nunca automático
        _tarefa_lower = tarefa.lower()
        if any(k in _tarefa_lower for k in ("overhaul", "rolamento", "desmontagem", "kit revisão", "kit revisao")):
            status = "Depende de análise preditiva"

        rows.append([tarefa[:60], tipo, prox_d, prox_h, status])

    _tbl(doc,
         ["Tarefa", "Tipo", "Próx. Data", "Próx. Horímetro", "Status"],
         rows, col_widths_cm=[6.0, 3.5, 2.5, 3.0, 5.0])


def _secao_relatorios(doc, dados: dict) -> None:
    _hd(doc, "8. Resumo dos Relatórios Técnicos")
    df = dados.get("relatorios")
    if df is None or df.empty:
        _txt(doc, "Nenhum relatório técnico publicado no período selecionado.", cor=_MUTED)
        return

    for _, row in df.iterrows():
        titulo  = str(row.get("Titulo",         "Relatório")).strip()
        tipo    = str(row.get("Tipo_Servico",   "")).strip()
        data    = str(row.get("Data_Relatorio", "")).strip()[:10]
        sev     = str(row.get("Severidade",     "")).strip()
        achados = str(row.get("Achados",        "") or row.get("Descricao", "")).strip()
        rec     = str(row.get("Recomendacoes",  "")).strip()

        _hd(doc, f"  {titulo}", nivel=2)
        from docx.shared import Pt
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"Tipo: {tipo}  ·  Data: {data}  ·  Severidade: {sev or '—'}")
        run.font.size = Pt(8.5)
        run.font.color.rgb = _rgb(_MUTED)
        if achados:
            _txt(doc, f"Principais achados: {achados[:400]}", size_pt=9)
        if rec:
            _txt(doc, f"Recomendações: {rec[:300]}", size_pt=9)
        doc.add_paragraph()


def _secao_principais_pontos(doc, dados: dict) -> None:
    _hd(doc, "9. Principais Pontos a Mencionar")
    pontos = _gerar_pontos_tecnicos(dados)
    for pt in pontos:
        p = doc.add_paragraph(style="List Bullet")
        from docx.shared import Pt
        run = p.add_run(pt)
        run.font.size = Pt(9.5)


def _secao_alertas_chamados(doc, dados: dict) -> None:
    _hd(doc, "10. Alertas e Chamados")

    df_al = dados.get("alertas")
    if df_al is not None and not df_al.empty:
        _hd(doc, "Alertas", nivel=2)
        rows = []
        for _, row in df_al.iterrows():
            rows.append([
                str(row.get("Data",       "")).strip()[:10],
                str(row.get("Titulo",     "")).strip()[:70],
                str(row.get("Prioridade", "")).strip(),
                str(row.get("Status",     "")).strip(),
            ])
        _tbl(doc, ["Data", "Título", "Prioridade", "Status"], rows,
             col_widths_cm=[2.5, 8.0, 3.5, 4.0])

    df_ch = dados.get("chamados")
    if df_ch is not None and not df_ch.empty:
        _hd(doc, "Chamados Técnicos", nivel=2)
        rows = []
        for _, row in df_ch.iterrows():
            rows.append([
                str(row.get("Data_Abertura","")).strip()[:10],
                str(row.get("Titulo",       "")).strip()[:70],
                str(row.get("Prioridade",   "")).strip(),
                str(row.get("Status",       "")).strip(),
            ])
        _tbl(doc, ["Abertura", "Título", "Prioridade", "Status"], rows,
             col_widths_cm=[2.5, 8.0, 3.5, 4.0])


def _secao_recomendacoes(doc, dados: dict) -> None:
    _hd(doc, "11. Recomendações Técnicas")
    df_rel = dados.get("relatorios")
    rows   = []

    if df_rel is not None and not df_rel.empty:
        for _, row in df_rel.iterrows():
            rec  = str(row.get("Recomendacoes", "")).strip()
            tipo = str(row.get("Tipo_Servico",  "")).strip()
            data = str(row.get("Data_Relatorio","")).strip()[:10]
            if rec:
                rows.append([rec[:100], tipo, "—", data, "—"])

    ativo = dados.get("ativo", {})
    det   = str(ativo.get("Detalhes", "")).strip()
    if det:
        rows.append([det[:100], "Avaliação do ativo", "Alta", "—", "Pendente"])

    if not rows:
        _txt(doc, "Nenhuma recomendação técnica registrada no período.", cor=_MUTED)
        return

    _tbl(doc,
         ["Recomendação", "Origem", "Prioridade", "Data", "Status"],
         rows, col_widths_cm=[7.0, 3.5, 2.5, 2.5, 2.5])


def _secao_recomendacoes_condicao(doc, dados: dict) -> None:
    _hd(doc, "12. Recomendações por Condição")
    _txt(doc,
         "As recomendações por condição NÃO são automáticas por horímetro. Elas dependem da "
         "saúde real do ativo, relatórios preditivos, histórico operacional e avaliação técnica Pred.IO.",
         size_pt=9, italic=True, color=_MUTED)

    itens = [
        ("Overhaul",               "Não indicado no momento — depende de análise preditiva"),
        ("Desmontagem do compressor","Não indicado no momento — depende de avaliação técnica"),
        ("Kit revisão",            "Não indicado no momento — depende de análise preditiva"),
        ("Troca de rolamento",     "Não indicado no momento — depende de análise de vibração"),
        ("Intervenção pesada",     "Em avaliação Pred.IO — monitorar tendência"),
    ]
    rows = [[it[0], it[1]] for it in itens]
    _tbl(doc, ["Item", "Status / Observação"], rows, col_widths_cm=[5.5, 12.5])


def _secao_conclusao(doc, dados: dict) -> None:
    _hd(doc, "13. Conclusão")
    texto = _gerar_texto_conclusao(dados)
    _txt(doc, texto, size_pt=10)
    doc.add_paragraph()
    _txt(doc,
         "Este relatório foi gerado automaticamente pelo Portal Pred.IO e deve ser revisado "
         "pela equipe técnica antes do envio ao cliente. O relatório não substitui a avaliação "
         "técnica presencial quando esta for necessária.",
         size_pt=8.5, italic=True, color=_MUTED)
    doc.add_paragraph()
    _txt(doc, "Fonte: Pred.IO  ·  Documento editável para revisão técnica", size_pt=8, color=_MUTED)


def _rodape_doc(doc) -> None:
    pass  # rodapé de página não alterado — python-docx suporta mas é opcional


# ═══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO DE TEXTOS AUTOMÁTICOS
# ═══════════════════════════════════════════════════════════════════════════════
def _gerar_texto_resumo(dados: dict) -> str:
    ativo        = dados.get("ativo", {})
    nome         = str(ativo.get("Tag", "")).strip() or str(ativo.get("Nome", "")).strip() or "Ativo"
    status       = str(ativo.get("Status", "Indefinido")).strip()
    score_raw    = str(ativo.get("Score", "—")).strip()
    cliente      = dados.get("cliente_nome", "")
    periodo_ini  = dados.get("periodo_inicio", "")
    periodo_fim  = dados.get("periodo_fim", "")

    n_rel  = 0 if dados.get("relatorios") is None or dados["relatorios"].empty else len(dados["relatorios"])
    n_cham = 0 if dados.get("chamados")   is None or dados["chamados"].empty   else len(dados["chamados"])
    n_al   = 0 if dados.get("alertas")    is None or dados["alertas"].empty    else len(dados["alertas"])

    periodo_str = f"de {periodo_ini} a {periodo_fim}" if periodo_ini and periodo_fim else "no período analisado"

    texto = (
        f"O ativo {nome} do cliente {cliente} apresenta status {status}, "
        f"com score de saúde {score_raw}/100. "
        f"No período analisado ({periodo_str}), foram registrados "
        f"{n_rel} relatório(s) técnico(s), {n_al} alerta(s) e {n_cham} chamado(s) técnico(s). "
    )

    # Reforçar regras de negócio
    texto += (
        "Não há indicação automática de overhaul, troca de rolamento, desmontagem ou "
        "kit revisão com base exclusiva em horímetro. "
        "Essas decisões dependem da evolução das análises preditivas e da avaliação técnica Pred.IO. "
        "Fonte: Pred.IO."
    )
    return texto


def _gerar_pontos_resumo(dados: dict) -> list[str]:
    ativo  = dados.get("ativo", {})
    nome   = str(ativo.get("Tag", "")).strip() or "Ativo"
    status = str(ativo.get("Status", "")).strip()
    score  = str(ativo.get("Score",  "—")).strip()
    pontos = [
        f"{nome} apresenta status {status} com score {score}/100.",
        "Não há indicação automática de overhaul ou troca de rolamento por horímetro.",
        "Decisões de revisão geral dependem de análise preditiva e avaliação técnica.",
        "Recomenda-se manter o plano de análise de óleo conforme programado.",
    ]
    df_al = dados.get("alertas")
    if df_al is not None and not df_al.empty:
        n_crit = 0
        if "Prioridade" in df_al.columns:
            n_crit = int(df_al["Prioridade"].str.lower().str.contains("crítica|urgente", na=False).sum())
        if n_crit > 0:
            pontos.append(f"{n_crit} alerta(s) crítico(s) registrado(s) no período — atenção recomendada.")
    return pontos


def _gerar_pontos_tecnicos(dados: dict) -> list[str]:
    ativo  = dados.get("ativo", {})
    nome   = str(ativo.get("Tag", "")).strip() or "Ativo"
    status = str(ativo.get("Status", "")).strip()
    pontos = [
        f"{nome} permanece em status {status}.",
        "Overhaul depende de avaliação por condição — não por horímetro.",
        "Troca de rolamento indicada apenas se análise de vibração confirmar necessidade.",
        "Análise de óleo deve ser mantida conforme plano vigente.",
        "Termografia não apresentou ponto crítico no período (confirmar com dados reais).",
        "Fonte: Pred.IO.",
    ]
    return pontos


def _gerar_texto_conclusao(dados: dict) -> str:
    ativo  = dados.get("ativo", {})
    nome   = str(ativo.get("Tag", "")).strip() or str(ativo.get("Nome", "")).strip() or "Ativo"
    status = str(ativo.get("Status", "Indefinido")).strip()
    score  = str(ativo.get("Score",  "—")).strip()

    return (
        f"O ativo {nome} encerra o período analisado com status {status} e score {score}/100. "
        "As principais ações recomendadas são: manter o plano de manutenção preventiva vigente, "
        "acompanhar a evolução do score nas próximas coletas, "
        "e aguardar avaliação técnica da equipe Pred.IO antes de qualquer intervenção pesada. "
        "Não há indicação automática de overhaul ou troca de rolamento com base em horímetro. "
        "Fonte: Pred.IO."
    )


# ── Compat: aceita ambos 'cor' e 'color' como keyword ─────────────────────────
_orig_txt = _txt
def _txt(doc, texto: str, size_pt: int = 10, bold: bool = False,
         color: str = "2D3748", cor: str = "", italic: bool = False) -> None:
    _orig_txt(doc, texto, size_pt=size_pt, bold=bold,
              color=cor if cor else color, italic=italic)
