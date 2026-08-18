"""Resumo Executivo por Período — Portal Pred.IO.

Consolida relatórios, manutenções, alertas, chamados, GUT e recomendações
de um cliente (e, opcionalmente, de um ativo) num resumo pronto para
reuniões de gerência.

SEGURANÇA:
- cliente_id deve SEMPRE vir da sessão (current_client_id()) no Portal do
  Cliente, ou da seleção explícita do staff na Supervisão — nunca de
  input livre do cliente comum.
- modo "cliente" e "admin_cliente" NUNCA incluem rascunhos nem
  observações internas — usam sempre os mesmos getters staff=False já
  usados no restante do Portal do Cliente (get_technical_reports,
  get_maintenance_tasks, get_gut_summary).
- modo "interno_predio" pode incluir rascunhos e observações internas.
  Quem chama esta função é responsável por só usar "interno_predio"
  depois de validar is_staff() — esta função não lê st.session_state
  para permanecer pura e testável fora do contexto Streamlit.
- Nenhuma nota GUT autoriza overhaul, troca de rolamento ou parada de
  máquina automaticamente — este módulo só reaproveita os textos de
  gut.gut_acao_recomendada(), que já respeitam essa regra.
"""
from __future__ import annotations

import datetime as _dt
import re
import unicodedata
import pandas as pd

from gut import GUT_DISCLAIMER

TIPOS_RESUMO = ["Resumo para reunião", "Resumo técnico", "Resumo gerencial"]
MODOS = ("cliente", "admin_cliente", "interno_predio")

# ── Categorias de exibição de Tipo_Servico (agrupamento Cliente → Tipo) ───────
#
# Tipo_Servico NÃO é um enum fechado — o formulário manual da Supervisão usa
# uma lista fixa, mas a integração do App Relatórios grava valores livres
# (ex.: "Ordem de Assistência (OAT)") direto no campo. Estas categorias são
# só de EXIBIÇÃO/filtro — nunca reescrevem o Tipo_Servico real gravado.
# Definidas aqui (não em page_sv_relatorios.py) para serem a fonte única
# tanto pra tela de Relatórios da Supervisão quanto pro filtro de tipo do
# Resumo Executivo, evitando duas listas de aliases divergindo com o tempo.

CATEGORIAS_RELATORIO = ["Vibração", "Termografia", "OAT", "Ordem de Serviço",
                        "Análise de Óleo", "Alinhamento a Laser",
                        "Relatório Executivo", "Outros"]

# Ordem de inserção importa: categoria_relatorio() retorna a primeira que
# bater — "OAT" precisa vir ANTES de "Ordem de Serviço" para que relatórios
# de Ordem de Assistência Técnica (OAT) tenham categoria própria em vez de
# caírem no balaio genérico de "Ordem de Serviço".
_ALIASES_CATEGORIA = {
    "Vibração":            ["vibracao", "vibração"],
    "Termografia":         ["termografia", "termografico", "térmico", "termico"],
    "OAT":                 ["oat", "ordem de assistencia", "ordem de assistência",
                            "relatorio oat", "relatório oat"],
    "Ordem de Serviço":    ["ordem de servico", "ordem de serviço", " os "],
    "Análise de Óleo":     ["oleo", "óleo"],
    "Alinhamento a Laser": ["alinhamento"],
    "Relatório Executivo": ["executivo"],
}


def _norm_tipo(s: str) -> str:
    n = unicodedata.normalize("NFD", str(s or "").lower())
    n = re.sub(r"[̀-ͯ]", "", n)
    return f" {n.strip()} "


def categoria_relatorio(tipo_servico: str) -> str:
    """Mapeia um Tipo_Servico real (livre) para uma categoria de exibição
    fixa — nunca inventa categoria fora de CATEGORIAS_RELATORIO, cai em
    "Outros" quando não reconhece."""
    t = _norm_tipo(tipo_servico)
    for categoria, aliases in _ALIASES_CATEGORIA.items():
        if any(a in t for a in aliases):
            return categoria
    return "Outros"

TEXTO_OBRIGATORIO = (
    "Este resumo consolida informações disponíveis no Portal Pred.IO no período "
    "selecionado e tem objetivo de apoio à gestão. Decisões críticas devem ser "
    "validadas pela equipe técnica Pred.IO."
)

PERIODOS_PRESET = {
    "7d":  ("Últimos 7 dias", 7),
    "30d": ("Últimos 30 dias", 30),
    "90d": ("Últimos 90 dias", 90),
}


# ═══════════════════════════════════════════════════════════════════════════════
# PERÍODO
# ═══════════════════════════════════════════════════════════════════════════════

def resolver_periodo(preset: str, custom_ini: _dt.date | None = None,
                     custom_fim: _dt.date | None = None) -> tuple:
    """Retorna (data_inicio, data_fim) a partir do preset selecionado.
    Preset padrão (desconhecido/ausente): últimos 30 dias."""
    hoje = _dt.date.today()
    if preset == "este_mes":
        return hoje.replace(day=1), hoje
    if preset == "mes_anterior":
        primeiro_atual   = hoje.replace(day=1)
        ultimo_anterior  = primeiro_atual - _dt.timedelta(days=1)
        return ultimo_anterior.replace(day=1), ultimo_anterior
    if preset == "custom" and custom_ini and custom_fim:
        return custom_ini, custom_fim
    _, dias = PERIODOS_PRESET.get(preset, PERIODOS_PRESET["30d"])
    return hoje - _dt.timedelta(days=dias), hoje


def _filtrar_periodo(df: pd.DataFrame, col_data: str, ini: _dt.date, fim: _dt.date) -> pd.DataFrame:
    if df is None or df.empty or col_data not in df.columns:
        return pd.DataFrame()
    d = df.copy()
    d["_dt"] = pd.to_datetime(d[col_data].astype(str), dayfirst=True, errors="coerce")
    dt_ini = pd.Timestamp(ini)
    dt_fim = pd.Timestamp(fim) + pd.Timedelta(hours=23, minutes=59, seconds=59)
    d = d[(d["_dt"] >= dt_ini) & (d["_dt"] <= dt_fim)]
    return d.drop(columns=["_dt"]).reset_index(drop=True)


# ═══════════════════════════════════════════════════════════════════════════════
# COLETA DE DADOS
# ═══════════════════════════════════════════════════════════════════════════════

def _coletar_dados(cliente_id: str, ativo_id: str, ini: _dt.date, fim: _dt.date,
                   modo: str, incluir: dict, tipo_servico: str = "") -> dict:
    from sheets import (
        get_technical_reports, get_maintenance_tasks, get_maintenance_executions,
        get_alertas_sv, get_chamados_v2, get_gut_summary, get_all_ativos_sv,
    )

    staff_mode = modo == "interno_predio"
    dados: dict = {}

    # ── Ativos monitorados ──────────────────────────────────────────────────
    df_ativos = pd.DataFrame()
    try:
        df_ativos = get_all_ativos_sv()
        if not df_ativos.empty and "Client_Id" in df_ativos.columns:
            df_ativos = df_ativos[
                df_ativos["Client_Id"].astype(str).str.strip().str.lower() == cliente_id.strip().lower()
            ]
            if ativo_id:
                df_ativos = df_ativos[df_ativos["Id"].astype(str).str.strip() == ativo_id.strip()]
    except Exception:
        df_ativos = pd.DataFrame()
    dados["ativos"] = df_ativos

    # ── Relatórios publicados (ou todos, em modo interno) ───────────────────
    df_rel = pd.DataFrame()
    if incluir.get("relatorios", True):
        try:
            df_rel = get_technical_reports(client_id=cliente_id, ativo_id=ativo_id, staff=staff_mode)
            df_rel = _filtrar_periodo(df_rel, "Data_Relatorio", ini, fim)
            if tipo_servico and not df_rel.empty:
                df_rel = df_rel[df_rel["Tipo_Servico"].apply(categoria_relatorio) == tipo_servico]
            if not staff_mode and "Obs_Interna" in df_rel.columns:
                df_rel = df_rel.drop(columns=["Obs_Interna"])
        except Exception:
            df_rel = pd.DataFrame()
    dados["relatorios"] = df_rel

    # ── Manutenções executadas no período ───────────────────────────────────
    df_mex = pd.DataFrame()
    if incluir.get("manutencoes", True):
        try:
            df_mex = get_maintenance_executions(client_id=cliente_id, ativo_id=ativo_id, limit=300)
            df_mex = _filtrar_periodo(df_mex, "Executado_Em", ini, fim)
            if not staff_mode and "Obs_Interna" in df_mex.columns:
                df_mex = df_mex.drop(columns=["Obs_Interna"])
        except Exception:
            df_mex = pd.DataFrame()
    dados["manutencoes_executadas"] = df_mex

    # ── Manutenções pendentes/vencidas (situação atual, não é por período) ──
    df_mpend = pd.DataFrame()
    if incluir.get("manutencoes", True):
        try:
            df_mpend = get_maintenance_tasks(client_id=cliente_id, ativo_id=ativo_id, staff=staff_mode)
            if not df_mpend.empty and "Status" in df_mpend.columns:
                df_mpend = df_mpend[~df_mpend["Status"].str.lower().str.contains("conclu|arquiv", na=False)]
        except Exception:
            df_mpend = pd.DataFrame()
    dados["manutencoes_pendentes"] = df_mpend

    # ── Alertas do período ───────────────────────────────────────────────────
    df_al = pd.DataFrame()
    if incluir.get("alertas", True):
        try:
            df_al = get_alertas_sv(cliente_id)
            if not df_al.empty and ativo_id and "Ativo_Id" in df_al.columns:
                df_al = df_al[df_al["Ativo_Id"].astype(str).str.strip() == ativo_id.strip()]
            df_al = _filtrar_periodo(df_al, "Criado_Em", ini, fim)
        except Exception:
            df_al = pd.DataFrame()
    dados["alertas"] = df_al

    # ── Chamados do período ─────────────────────────────────────────────────
    df_cham = pd.DataFrame()
    if incluir.get("chamados", True):
        try:
            df_cham = get_chamados_v2(client_id=cliente_id, ativo_id=ativo_id)
            df_cham = _filtrar_periodo(df_cham, "Aberto_Em", ini, fim)
        except Exception:
            df_cham = pd.DataFrame()
    dados["chamados"] = df_cham

    # ── GUT / prioridades — sempre client-safe (get_gut_summary já filtra
    #    com staff=False internamente, mesmo em modo interno_predio) ─────────
    gut_itens: list = []
    if incluir.get("gut", True):
        try:
            gut_itens = get_gut_summary(cliente_id)
            if ativo_id:
                gut_itens = [i for i in gut_itens if i.get("ativo_id", "") == ativo_id]

            def _no_periodo(item: dict) -> bool:
                d = pd.to_datetime(item.get("created_at", ""), dayfirst=True, errors="coerce")
                if pd.isna(d):
                    return True  # sem data legível — mantém para não perder item relevante
                return pd.Timestamp(ini) <= d <= pd.Timestamp(fim) + pd.Timedelta(hours=23, minutes=59)

            gut_itens = [i for i in gut_itens if _no_periodo(i)]
        except Exception:
            gut_itens = []
    dados["gut_itens"] = gut_itens

    return dados


# ═══════════════════════════════════════════════════════════════════════════════
# ANALYTICS — indicadores, gráficos, pontos para gerência, ações, timeline
# ═══════════════════════════════════════════════════════════════════════════════

def _sem_acento(s: str) -> str:
    import unicodedata
    return "".join(
        c for c in unicodedata.normalize("NFD", str(s).lower().strip())
        if unicodedata.category(c) != "Mn"
    )


def _status_ativo_norm(raw: str) -> str:
    s = _sem_acento(raw)
    if s in ("bom", "verde", "normal", "ok"):
        return "Bom"
    if s in ("atencao", "amarelo", "alerta"):
        return "Atenção"
    if s in ("critico", "critica", "vermelho"):
        return "Crítico"
    if s == "urgente":
        return "Urgente"
    return "—"


def compute_indicadores(dados: dict) -> dict:
    """Indicadores gerenciais em número — para os cards do resumo."""
    df_ativos = dados.get("ativos")
    df_rel    = dados.get("relatorios")
    df_mex    = dados.get("manutencoes_executadas")
    df_mpend  = dados.get("manutencoes_pendentes")
    df_al     = dados.get("alertas")
    df_cham   = dados.get("chamados")
    gut_itens = dados.get("gut_itens", [])

    n_ativos = len(df_ativos) if df_ativos is not None else 0
    score_medio = None
    n_atencao = n_critico = n_urgente = 0
    if df_ativos is not None and not df_ativos.empty:
        if "Score" in df_ativos.columns:
            scores = pd.to_numeric(df_ativos["Score"], errors="coerce").dropna()
            if len(scores):
                score_medio = round(scores.mean(), 1)
        if "Status" in df_ativos.columns:
            st_norm = df_ativos["Status"].astype(str).apply(_status_ativo_norm)
            n_atencao = int((st_norm == "Atenção").sum())
            n_critico = int((st_norm == "Crítico").sum())
            n_urgente = int((st_norm == "Urgente").sum())

    n_manut_vencidas = 0
    if df_mpend is not None and not df_mpend.empty and "Status" in df_mpend.columns:
        n_manut_vencidas = int(df_mpend["Status"].astype(str).str.lower().str.contains("vencid|atraso").sum())

    n_alertas_criticos = 0
    if df_al is not None and not df_al.empty and "Prioridade" in df_al.columns:
        n_alertas_criticos = int(df_al["Prioridade"].astype(str).str.lower().isin(["crítica", "critica", "urgente"]).sum())

    n_chamados_abertos = 0
    if df_cham is not None and not df_cham.empty and "Status" in df_cham.columns:
        n_chamados_abertos = int((df_cham["Status"].astype(str).str.lower() != "concluído").sum())

    n_gut_criticos = sum(1 for i in gut_itens if i.get("prioridade") == "Crítica")

    return {
        "ativos_monitorados": n_ativos,
        "saude_media": score_medio,
        "ativos_atencao": n_atencao,
        "ativos_criticos": n_critico + n_urgente,
        "manutencoes_vencidas": n_manut_vencidas,
        "itens_gut_criticos": n_gut_criticos,
        "alertas_criticos": n_alertas_criticos,
        "chamados_abertos": n_chamados_abertos,
        "relatorios_publicados": len(df_rel) if df_rel is not None else 0,
        "manutencoes_executadas": len(df_mex) if df_mex is not None else 0,
    }


def compute_chart_data(dados: dict) -> dict:
    """Dados agregados prontos para gráfico (dict de labels->valor, ou lista
    de pontos). Chaves com dado insuficiente ficam com valor "falsy" (dict
    vazio / lista vazia) — quem renderiza deve mostrar estado amigável
    nesses casos, nunca um gráfico vazio."""
    df_ativos = dados.get("ativos")
    df_rel    = dados.get("relatorios")
    df_mpend  = dados.get("manutencoes_pendentes")
    df_mex    = dados.get("manutencoes_executadas")
    gut_itens = dados.get("gut_itens", [])

    # ── 1. Saúde dos ativos ─────────────────────────────────────────────────
    saude_ativos: dict = {}
    if df_ativos is not None and not df_ativos.empty and "Status" in df_ativos.columns:
        st_norm = df_ativos["Status"].astype(str).apply(_status_ativo_norm)
        for label in ("Bom", "Atenção", "Crítico", "Urgente"):
            n = int((st_norm == label).sum())
            if n:
                saude_ativos[label] = n

    # ── 2. Evolução do score de saúde (aproximada por Score_Impacto dos
    #      relatórios publicados no período, ancorada na saúde média atual;
    #      só é exibida se houver pelo menos 2 pontos de data distintos) ────
    score_evolucao: list = []
    if df_rel is not None and not df_rel.empty and "Score_Impacto" in df_rel.columns \
            and "Data_Relatorio" in df_rel.columns and df_ativos is not None and not df_ativos.empty:
        try:
            tmp = df_rel[["Data_Relatorio", "Score_Impacto"]].copy()
            tmp["_dt"]  = pd.to_datetime(tmp["Data_Relatorio"], dayfirst=True, errors="coerce")
            tmp["_imp"] = pd.to_numeric(tmp["Score_Impacto"], errors="coerce").fillna(0)
            tmp = tmp.dropna(subset=["_dt"]).sort_values("_dt")
            if tmp["_dt"].nunique() >= 2:
                scores_atuais = pd.to_numeric(df_ativos.get("Score", pd.Series(dtype=float)), errors="coerce").dropna()
                score_final = float(scores_atuais.mean()) if len(scores_atuais) else 75.0
                # Caminha de trás para frente: score no ponto = final - soma dos impactos posteriores
                impactos_pos = tmp["_imp"][::-1].cumsum()[::-1]
                pontos = []
                for (_, row), impacto_acumulado in zip(tmp.iterrows(), impactos_pos):
                    valor = max(0, min(100, score_final - impacto_acumulado + row["_imp"]))
                    pontos.append((row["_dt"].strftime("%d/%m"), round(valor, 1)))
                pontos.append((_dt.date.today().strftime("%d/%m"), round(score_final, 1)))
                score_evolucao = pontos
        except Exception:
            score_evolucao = []

    # ── 3. GUT por prioridade ────────────────────────────────────────────────
    gut_prioridade: dict = {}
    for label in ("Baixa", "Moderada", "Alta", "Crítica"):
        n = sum(1 for i in gut_itens if i.get("prioridade") == label)
        if n:
            gut_prioridade[label] = n

    # ── 4. Manutenções por status ────────────────────────────────────────────
    manut_status: dict = {}
    if df_mpend is not None and not df_mpend.empty and "Status" in df_mpend.columns:
        s_norm = df_mpend["Status"].astype(str).apply(_sem_acento)
        tipo_norm = df_mpend.get("Tipo_Manutencao", pd.Series(dtype=str)).astype(str).apply(_sem_acento)
        n_condicao = int((tipo_norm == "condicao").sum())
        n_vencidas = int((s_norm.str.contains("vencid|atraso") & (tipo_norm != "condicao")).sum())
        n_proximas = int((s_norm.str.contains("proxim") & (tipo_norm != "condicao")).sum())
        n_em_dia   = int(len(df_mpend)) - n_condicao - n_vencidas - n_proximas
        for label, n in [("Em dia", max(n_em_dia, 0)), ("Próximas", n_proximas),
                         ("Vencidas", n_vencidas), ("Por condição", n_condicao)]:
            if n:
                manut_status[label] = n
    if df_mex is not None and not df_mex.empty:
        manut_status["Concluídas"] = len(df_mex)

    # ── 5. Top 5 ativos mais críticos (score baixo / maior GUT) ─────────────
    top5_criticos: list = []
    if df_ativos is not None and not df_ativos.empty:
        gut_por_ativo: dict = {}
        for i in gut_itens:
            aid = i.get("ativo_id", "")
            if aid and (aid not in gut_por_ativo or i["score"] > gut_por_ativo[aid]["score"]):
                gut_por_ativo[aid] = i
        id_col   = "Id" if "Id" in df_ativos.columns else df_ativos.columns[0]
        nome_col = "Tag" if "Tag" in df_ativos.columns else id_col
        linhas = []
        for _, row in df_ativos.iterrows():
            aid   = str(row.get(id_col, "")).strip()
            score = pd.to_numeric(row.get("Score", None), errors="coerce")
            score = float(score) if pd.notna(score) else None
            gut_i = gut_por_ativo.get(aid)
            linhas.append({
                "ativo": str(row.get(nome_col, aid)).strip() or aid,
                "score": score,
                "gut_score": gut_i["score"] if gut_i else None,
                "gut_prioridade": gut_i["prioridade"] if gut_i else "—",
            })
        linhas.sort(key=lambda r: (r["gut_score"] or 0, -(r["score"] or 100)), reverse=True)
        top5_criticos = linhas[:5]

    # ── 6. Relatórios por severidade ────────────────────────────────────────
    rel_severidade: dict = {}
    if df_rel is not None and not df_rel.empty and "Severidade" in df_rel.columns:
        for label in ("Normal", "Atenção", "Crítico", "Urgente"):
            n = int(df_rel["Severidade"].astype(str).str.strip().eq(label).sum())
            if n:
                rel_severidade[label] = n

    return {
        "saude_ativos":         saude_ativos,
        "score_evolucao":       score_evolucao,
        "gut_prioridade":       gut_prioridade,
        "manutencoes_status":   manut_status,
        "top5_criticos":        top5_criticos,
        "relatorios_severidade": rel_severidade,
    }


def detectar_recorrencias(dados: dict) -> list[str]:
    """Detecta o mesmo ativo+tipo de serviço aparecendo em 2+ relatórios
    publicados dentro do período — ex.: "Motor X apresentou recorrência de
    Análise de Vibração em 2 relatórios no período." Só usa metadados já
    presentes nos relatórios (Ativo_Id/Tipo_Servico), sem heurística de
    texto livre. Ordenado das recorrências mais frequentes para as menos."""
    df_rel = dados.get("relatorios")
    if df_rel is None or df_rel.empty \
            or "Ativo_Id" not in df_rel.columns or "Tipo_Servico" not in df_rel.columns:
        return []

    df_ativos = dados.get("ativos")
    nome_por_id: dict = {}
    if df_ativos is not None and not df_ativos.empty:
        id_col   = "Id"  if "Id"  in df_ativos.columns else df_ativos.columns[0]
        nome_col = "Tag" if "Tag" in df_ativos.columns else id_col
        for _, row in df_ativos.iterrows():
            nome_por_id[str(row.get(id_col, "")).strip()] = str(row.get(nome_col, "")).strip()

    tmp = df_rel.copy()
    tmp["_ativo"] = tmp["Ativo_Id"].astype(str).str.strip()
    tmp["_tipo"]  = tmp["Tipo_Servico"].astype(str).str.strip()
    tmp = tmp[(tmp["_ativo"] != "") & (tmp["_tipo"] != "")]

    grupos: list[tuple[int, str]] = []
    for (ativo_id, tipo), grupo in tmp.groupby(["_ativo", "_tipo"]):
        n = len(grupo)
        if n >= 2:
            nome = nome_por_id.get(ativo_id, "") or ativo_id
            grupos.append((n, f"{nome} apresentou recorrência de {tipo} em {n} relatórios no período."))
    grupos.sort(key=lambda t: t[0], reverse=True)
    return [frase for _, frase in grupos]


def principais_pontos_gerencia(dados: dict, indicadores: dict) -> list:
    """Até 5 pontos objetivos para levar à reunião — frases curtas, sem
    parágrafos longos. Recorrências entre relatórios (mesmo ativo+tipo
    aparecendo 2+ vezes) vêm primeiro — é um sinal mais forte que um único
    relatório isolado. A última linha (segurança/transparência) é sempre
    incluída quando há espaço, espelhando a regra de nunca indicar
    overhaul/rolamento automaticamente."""
    pontos: list = []
    df_ativos = dados.get("ativos")
    gut_itens = dados.get("gut_itens", [])
    df_rel    = dados.get("relatorios")

    # Recorrências entre relatórios — prioridade mais alta
    pontos.extend(detectar_recorrencias(dados))

    # Ativo com pior score fora de "Bom"
    if df_ativos is not None and not df_ativos.empty and "Score" in df_ativos.columns:
        tmp = df_ativos.copy()
        tmp["_score"] = pd.to_numeric(tmp["Score"], errors="coerce")
        tmp = tmp.dropna(subset=["_score"])
        tmp = tmp[tmp["Status"].astype(str).apply(_status_ativo_norm) != "Bom"] if "Status" in tmp.columns else tmp
        if not tmp.empty:
            pior = tmp.sort_values("_score").iloc[0]
            nome_col = "Tag" if "Tag" in tmp.columns else tmp.columns[0]
            status_lbl = _status_ativo_norm(pior.get("Status", "")) if "Status" in tmp.columns else "atenção"
            pontos.append(f"{pior.get(nome_col, 'Ativo')} permanece em {status_lbl.lower()}, com score {int(pior['_score'])}.")

    # Item GUT de maior prioridade
    top_gut = sorted(gut_itens, key=lambda i: i.get("score", 0), reverse=True)
    if top_gut and top_gut[0].get("prioridade") in ("Alta", "Crítica"):
        i = top_gut[0]
        pontos.append(f"{i.get('titulo','Item')} possui prioridade GUT {i.get('prioridade','').lower()}.")

    # Manutenções vencidas
    n_venc = indicadores.get("manutencoes_vencidas", 0)
    if n_venc:
        plural = "estão vencidas" if n_venc > 1 else "está vencida"
        pontos.append(f"{n_venc} manutenç{'ões' if n_venc > 1 else 'ão'} {plural}.")

    # Recomendação de relatório de maior severidade
    if df_rel is not None and not df_rel.empty:
        sev_rank = {"Urgente": 0, "Crítico": 1, "Atenção": 2, "Normal": 3}
        tmp = df_rel.copy()
        tmp["_r"] = tmp.get("Severidade", "Normal").map(sev_rank).fillna(9)
        pior_rel = tmp.sort_values("_r").iloc[0]
        rec = str(pior_rel.get("Recomendacoes", "")).strip()
        if rec and rec.lower() not in ("nan", ""):
            titulo = str(pior_rel.get("Titulo", "Relatório")).strip()
            pontos.append(f"O relatório «{titulo}» recomenda: {rec[:100]}.")

    # Linha fixa de transparência — sempre por último quando há espaço
    pontos.append("Não há indicação automática de overhaul ou troca de rolamento no período.")

    return pontos[:5]


_RANK_ORIGEM = {"manutencao_vencida": 3, "alerta_critico": 4}


def acoes_prioritarias(dados: dict, limite: int = 8) -> list:
    """Lista ordenada: 1) GUT crítica, 2) GUT alta, 3) manutenção vencida
    sem GUT, 4) alerta crítico sem GUT, 5) recomendação por condição.
    Cada item: {rank, acao, ativo_id, origem, prioridade, prazo, status}."""
    gut_itens = dados.get("gut_itens", [])
    df_mpend  = dados.get("manutencoes_pendentes")
    df_al     = dados.get("alertas")

    usados_ids = set()
    acoes: list = []

    for i in gut_itens:
        if i.get("prioridade") == "Crítica":
            acoes.append({"rank": 1, "acao": i.get("titulo", ""), "ativo_id": i.get("ativo_id", ""),
                          "origem": i.get("origem", ""), "prioridade": "Crítica",
                          "prazo": i.get("created_at", ""), "status": i.get("status", "")})
            usados_ids.add(i.get("id", ""))
    for i in gut_itens:
        if i.get("prioridade") == "Alta":
            acoes.append({"rank": 2, "acao": i.get("titulo", ""), "ativo_id": i.get("ativo_id", ""),
                          "origem": i.get("origem", ""), "prioridade": "Alta",
                          "prazo": i.get("created_at", ""), "status": i.get("status", "")})
            usados_ids.add(i.get("id", ""))

    if df_mpend is not None and not df_mpend.empty and "Status" in df_mpend.columns:
        venc = df_mpend[df_mpend["Status"].astype(str).str.lower().str.contains("vencid|atraso")]
        for _, row in venc.iterrows():
            if str(row.get("Id", "")) in usados_ids:
                continue
            acoes.append({"rank": 3, "acao": str(row.get("Nome_Tarefa", "Manutenção")).strip(),
                          "ativo_id": str(row.get("Ativo_Id", "")).strip(), "origem": "manutencao",
                          "prioridade": str(row.get("Prioridade", "")).strip(),
                          "prazo": str(row.get("Proxima_Execucao_Data", "")).strip(),
                          "status": str(row.get("Status", "")).strip()})

    if df_al is not None and not df_al.empty and "Prioridade" in df_al.columns:
        criticos = df_al[df_al["Prioridade"].astype(str).str.lower().isin(["crítica", "critica", "urgente"])]
        for _, row in criticos.iterrows():
            if str(row.get("Id", "")) in usados_ids:
                continue
            acoes.append({"rank": 4, "acao": str(row.get("Titulo", "Alerta")).strip(),
                          "ativo_id": str(row.get("Ativo_Id", "")).strip(), "origem": "alerta",
                          "prioridade": str(row.get("Prioridade", "")).strip(),
                          "prazo": str(row.get("Criado_Em", "")).strip(), "status": ""})

    for i in gut_itens:
        if i.get("subtipo") == "Condição" and i.get("id", "") not in usados_ids \
                and i.get("prioridade") not in ("Crítica", "Alta"):
            acoes.append({"rank": 5, "acao": i.get("titulo", ""), "ativo_id": i.get("ativo_id", ""),
                          "origem": "recomendacao_condicao", "prioridade": i.get("prioridade", ""),
                          "prazo": i.get("created_at", ""), "status": i.get("status", "")})

    acoes.sort(key=lambda a: a["rank"])
    return acoes[:limite]


def timeline_periodo(dados: dict, limite: int = 8) -> list:
    """Eventos relevantes do período, mais recentes primeiro — não é o
    histórico bruto, só os eventos que importam para a reunião."""
    eventos: list = []

    df_rel = dados.get("relatorios")
    if df_rel is not None and not df_rel.empty:
        for _, r in df_rel.iterrows():
            eventos.append((str(r.get("Data_Relatorio", "")).strip(), "📁 Relatório publicado",
                            str(r.get("Titulo", "")).strip()))

    df_mex = dados.get("manutencoes_executadas")
    if df_mex is not None and not df_mex.empty:
        for _, m in df_mex.iterrows():
            eventos.append((str(m.get("Executado_Em", "")).strip(), "🔧 Manutenção executada",
                            str(m.get("Descricao_Execucao", "")).strip()[:80]))

    df_al = dados.get("alertas")
    if df_al is not None and not df_al.empty and "Prioridade" in df_al.columns:
        relevantes = df_al[df_al["Prioridade"].astype(str).str.lower().isin(["crítica", "critica", "alta", "urgente"])]
        for _, a in relevantes.iterrows():
            eventos.append((str(a.get("Criado_Em", "")).strip(), "🔔 Alerta",
                            str(a.get("Titulo", "")).strip()))

    df_cham = dados.get("chamados")
    if df_cham is not None and not df_cham.empty and "Prioridade" in df_cham.columns:
        relevantes = df_cham[df_cham["Prioridade"].astype(str).str.lower().isin(["crítica", "critica", "alta", "urgente"])]
        for _, c in relevantes.iterrows():
            eventos.append((str(c.get("Aberto_Em", "")).strip(), "🎫 Chamado",
                            str(c.get("Titulo", "")).strip()))

    def _key(ev):
        d = pd.to_datetime(ev[0], dayfirst=True, errors="coerce")
        return d if pd.notna(d) else pd.Timestamp.min

    eventos.sort(key=_key, reverse=True)
    return eventos[:limite]


# ═══════════════════════════════════════════════════════════════════════════════
# MONTAGEM DO TEXTO
# ═══════════════════════════════════════════════════════════════════════════════

def _fmt_data(d: _dt.date) -> str:
    return d.strftime("%d/%m/%Y")


def _linha(rotulo: str, valor) -> str:
    return f"  • {rotulo}: {valor}"


def _montar_texto(cliente_nome: str, ativo_nome: str, periodo_inicio: _dt.date,
                  periodo_fim: _dt.date, tipo_resumo: str, modo: str, dados: dict) -> str:
    df_ativos = dados["ativos"]
    df_rel    = dados["relatorios"]
    df_mex    = dados["manutencoes_executadas"]
    df_mpend  = dados["manutencoes_pendentes"]
    df_al     = dados["alertas"]
    df_cham   = dados["chamados"]
    gut_itens = dados["gut_itens"]

    n_ativos    = len(df_ativos) if df_ativos is not None else 0
    score_medio = None
    if df_ativos is not None and not df_ativos.empty and "Score" in df_ativos.columns:
        scores = pd.to_numeric(df_ativos["Score"], errors="coerce").dropna()
        if len(scores):
            score_medio = round(scores.mean(), 1)

    n_manut_vencidas = 0
    n_manut_proximas = 0
    if df_mpend is not None and not df_mpend.empty and "Status" in df_mpend.columns:
        s = df_mpend["Status"].astype(str).str.lower()
        n_manut_vencidas = int(s.str.contains("vencid|atraso").sum())
        n_manut_proximas = int(s.str.contains("próxim|proxim").sum())

    criticos_gut = [i for i in gut_itens if i.get("prioridade") == "Crítica"]
    altos_gut    = [i for i in gut_itens if i.get("prioridade") == "Alta"]

    linhas: list[str] = []

    # ── Título / subtítulo ──────────────────────────────────────────────────
    linhas.append("RESUMO EXECUTIVO PRED.IO")
    subtitulo = f"Cliente: {cliente_nome}"
    if ativo_nome:
        subtitulo += f"  ·  Ativo: {ativo_nome}"
    subtitulo += f"  ·  Período: {_fmt_data(periodo_inicio)} a {_fmt_data(periodo_fim)}"
    subtitulo += f"  ·  Tipo: {tipo_resumo}"
    linhas.append(subtitulo)
    linhas.append("")
    linhas.append(TEXTO_OBRIGATORIO)
    linhas.append("")

    # ── 1. Sumário executivo ────────────────────────────────────────────────
    linhas.append("1. SUMÁRIO EXECUTIVO")
    linhas.append(
        f"  No período de {_fmt_data(periodo_inicio)} a {_fmt_data(periodo_fim)}, "
        f"{n_ativos} ativo(s) monitorado(s)"
        + (f", com saúde média de {score_medio}/100" if score_medio is not None else "")
        + f". Foram registrados {len(df_rel)} relatório(s), {len(df_al)} alerta(s) e "
        f"{len(df_cham)} chamado(s). {len(criticos_gut)} item(ns) em prioridade GUT crítica."
    )
    linhas.append("")

    # ── 2. Indicadores principais ───────────────────────────────────────────
    linhas.append("2. INDICADORES PRINCIPAIS")
    linhas.append(_linha("Ativos monitorados", n_ativos))
    linhas.append(_linha("Saúde média", f"{score_medio}/100" if score_medio is not None else "—"))
    linhas.append(_linha("Relatórios publicados no período", len(df_rel)))
    linhas.append(_linha("Manutenções executadas no período", len(df_mex)))
    linhas.append(_linha("Manutenções vencidas (situação atual)", n_manut_vencidas))
    linhas.append(_linha("Manutenções próximas do vencimento", n_manut_proximas))
    linhas.append(_linha("Alertas no período", len(df_al)))
    linhas.append(_linha("Chamados no período", len(df_cham)))
    linhas.append(_linha("Itens GUT críticos", len(criticos_gut)))
    linhas.append(_linha("Itens GUT alta prioridade", len(altos_gut)))
    linhas.append("")

    # ── 3. Relatórios analisados ────────────────────────────────────────────
    linhas.append("3. RELATÓRIOS ANALISADOS")
    if df_rel is None or df_rel.empty:
        linhas.append("  Nenhum relatório no período selecionado.")
    else:
        for _, r in df_rel.head(10).iterrows():
            linhas.append(
                f"  • [{str(r.get('Data_Relatorio','')).strip()}] "
                f"{str(r.get('Titulo','')).strip()} — "
                f"{str(r.get('Tipo_Servico','')).strip()} "
                f"(Severidade: {str(r.get('Severidade','Normal')).strip()})"
            )
    linhas.append("")

    # ── 4. Principais pontos técnicos ───────────────────────────────────────
    linhas.append("4. PRINCIPAIS PONTOS TÉCNICOS")
    achados = []
    if df_rel is not None and not df_rel.empty:
        for _, r in df_rel.head(6).iterrows():
            # Resumo Técnico (resumo_tecnico) é a fonte prioritária quando
            # preenchido; Diagnóstico só entra como fallback quando o
            # relatório ainda não tem resumo técnico.
            resumo = str(r.get("Resumo", "")).strip()
            if not resumo or resumo.lower() == "nan":
                resumo = str(r.get("Diagnostico", "")).strip()
            if resumo and resumo.lower() not in ("nan", ""):
                achados.append(f"  • {str(r.get('Titulo','Relatório')).strip()}: {resumo[:220]}")
    if achados:
        linhas.extend(achados)
    else:
        linhas.append("  Nenhum achado técnico relevante registrado no período.")
    linhas.append("")

    # ── 5. Manutenções e ações realizadas ───────────────────────────────────
    linhas.append("5. MANUTENÇÕES E AÇÕES REALIZADAS")
    if df_mex is None or df_mex.empty:
        linhas.append("  Nenhuma execução de manutenção registrada no período.")
    else:
        for _, m in df_mex.head(10).iterrows():
            linhas.append(
                f"  • [{str(m.get('Executado_Em','')).strip()}] "
                f"{str(m.get('Descricao_Execucao','Execução de manutenção')).strip()[:160]} "
                f"(Responsável: {str(m.get('Responsavel','—')).strip()})"
            )
    linhas.append("")

    # ── 6. Pendências e próximas ações ──────────────────────────────────────
    linhas.append("6. PENDÊNCIAS E PRÓXIMAS AÇÕES")
    if df_mpend is None or df_mpend.empty:
        linhas.append("  Nenhuma manutenção pendente no momento.")
    else:
        for _, t in df_mpend.head(10).iterrows():
            linhas.append(
                f"  • {str(t.get('Nome_Tarefa','Tarefa')).strip()} — "
                f"Status: {str(t.get('Status','')).strip()} "
                f"(Próxima execução: {str(t.get('Proxima_Execucao_Data','—')).strip()})"
            )
    linhas.append("")

    # ── 7. GUT e prioridades críticas ───────────────────────────────────────
    linhas.append("7. GUT E PRIORIDADES CRÍTICAS")
    linhas.append(f"  {GUT_DISCLAIMER}")
    top_gut = sorted(gut_itens, key=lambda i: i.get("score", 0), reverse=True)[:10]
    if not top_gut:
        linhas.append("  Nenhum item com GUT calculado no período.")
    else:
        for i in top_gut:
            linhas.append(
                f"  • [{i.get('prioridade','')} · score {i.get('score','')}] "
                f"{i.get('titulo','')} ({i.get('origem','')}) — {i.get('acao_recomendada','')}"
            )
    linhas.append("")

    # ── 8. Alertas e chamados relevantes ────────────────────────────────────
    linhas.append("8. ALERTAS E CHAMADOS RELEVANTES")
    if df_al is not None and not df_al.empty:
        linhas.append("  Alertas:")
        for _, a in df_al.head(8).iterrows():
            linhas.append(f"    • [{str(a.get('Prioridade','')).strip()}] {str(a.get('Titulo','')).strip()}")
    else:
        linhas.append("  Nenhum alerta no período.")
    if df_cham is not None and not df_cham.empty:
        linhas.append("  Chamados:")
        for _, c in df_cham.head(8).iterrows():
            linhas.append(
                f"    • [{str(c.get('Status','')).strip()}] {str(c.get('Titulo','')).strip()} "
                f"(Prioridade: {str(c.get('Prioridade','')).strip()})"
            )
    else:
        linhas.append("  Nenhum chamado no período.")
    linhas.append("")

    # ── 9. Recomendações Pred.IO ────────────────────────────────────────────
    linhas.append("9. RECOMENDAÇÕES PRED.IO")
    recs = []
    for i in top_gut[:6]:
        acao = i.get("acao_recomendada", "")
        if acao and acao not in recs:
            recs.append(acao)
    if df_rel is not None and not df_rel.empty:
        for _, r in df_rel.head(6).iterrows():
            rec = str(r.get("Recomendacoes", "")).strip()
            if rec and rec.lower() not in ("nan", "") and rec not in recs:
                recs.append(rec[:200])
    if recs:
        for rec in recs[:10]:
            linhas.append(f"  • {rec}")
    else:
        linhas.append("  Nenhuma recomendação adicional no período.")
    linhas.append(
        "  Observação: overhaul, troca de rolamento e parada de máquina nunca são "
        "automáticos — dependem sempre de análise preditiva e avaliação técnica Pred.IO."
    )
    linhas.append("")

    # ── 10. Conclusão para reunião ──────────────────────────────────────────
    linhas.append("10. CONCLUSÃO PARA REUNIÃO")
    pontos_reuniao = []
    if n_manut_vencidas:
        pontos_reuniao.append(f"{n_manut_vencidas} manutenção(ões) vencida(s) requer(em) atenção imediata")
    if criticos_gut:
        pontos_reuniao.append(f"{len(criticos_gut)} item(ns) em prioridade GUT crítica")
    if df_cham is not None and not df_cham.empty:
        abertos = df_cham[df_cham.get("Status", "").astype(str).str.lower() != "concluído"] if "Status" in df_cham.columns else df_cham
        if len(abertos):
            pontos_reuniao.append(f"{len(abertos)} chamado(s) em aberto no período")
    if not pontos_reuniao:
        pontos_reuniao.append("Nenhum ponto crítico pendente identificado no período — operação estável")
    linhas.append("  Pontos para levar à reunião com a gerência:")
    for p in pontos_reuniao:
        linhas.append(f"  • {p}")
    linhas.append("")
    linhas.append("  Fonte: Pred.IO")

    return "\n".join(linhas)


# ═══════════════════════════════════════════════════════════════════════════════
# SERVIÇO PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def generate_executive_summary(
    *,
    usuario_id: str,
    cliente_id: str,
    cliente_nome: str,
    ativo_id: str = "",
    ativo_nome: str = "",
    periodo_inicio: _dt.date,
    periodo_fim: _dt.date,
    tipo_resumo: str = "Resumo para reunião",
    modo: str = "cliente",
    incluir: dict | None = None,
    salvar: bool = True,
    tipo_servico: str = "",
) -> dict:
    """Gera (e, por padrão, salva) o resumo executivo consolidado do período.

    Retorna dict:
      ok, titulo, texto, dados (DataFrames/listas usados — auditoria),
      cliente_id, ativo_id, periodo_inicio, periodo_fim, modo, tipo_resumo,
      report_ids_usados (Ids dos TechnicalReports publicados que entraram
      neste resumo — auditoria de fonte), summary_id (se salvar=True e a
      gravação funcionar).

    tipo_servico: categoria de exibição (uma de CATEGORIAS_RELATORIO,
    ex.: "Vibração") — se informada, só relatórios cujo Tipo_Servico real
    caia nessa categoria (via categoria_relatorio()) entram no resumo.
    "" (padrão) = todos os tipos.

    SEGURANÇA: cliente_id deve vir sempre de current_client_id() (Portal do
    Cliente) ou de uma seleção explícita do staff já autenticado como
    is_staff() (Supervisão) — nunca de input livre do cliente comum. Esta
    função não valida perfil porque não tem acesso à sessão Streamlit;
    quem chama (UI) é responsável por essa checagem antes de invocá-la
    com modo != "cliente".
    """
    if not cliente_id:
        return {"ok": False, "erro": "cliente_id obrigatório."}
    if modo not in MODOS:
        modo = "cliente"
    if periodo_fim < periodo_inicio:
        return {"ok": False, "erro": "Período inválido: fim anterior ao início."}

    incluir = incluir or {"relatorios": True, "manutencoes": True,
                          "alertas": True, "chamados": True, "gut": True}

    dados = _coletar_dados(cliente_id, ativo_id, periodo_inicio, periodo_fim, modo, incluir,
                           tipo_servico=tipo_servico)

    titulo = f"Resumo Executivo Pred.IO — {cliente_nome}"
    if ativo_nome:
        titulo += f" — {ativo_nome}"

    texto = _montar_texto(
        cliente_nome=cliente_nome, ativo_nome=ativo_nome,
        periodo_inicio=periodo_inicio, periodo_fim=periodo_fim,
        tipo_resumo=tipo_resumo, modo=modo, dados=dados,
    )

    indicadores = compute_indicadores(dados)
    charts      = compute_chart_data(dados)
    pontos      = principais_pontos_gerencia(dados, indicadores)
    acoes       = acoes_prioritarias(dados)
    timeline    = timeline_periodo(dados)

    # Auditoria de fonte — quais relatórios publicados entraram neste resumo,
    # pra "Relatórios utilizados: X" na UI e pra saber de onde veio cada dado.
    df_rel_usados = dados.get("relatorios")
    report_ids_usados = (
        df_rel_usados["Id"].astype(str).str.strip().tolist()
        if df_rel_usados is not None and not df_rel_usados.empty and "Id" in df_rel_usados.columns
        else []
    )

    resultado = {
        "ok": True, "titulo": titulo, "texto": texto, "dados": dados,
        "cliente_id": cliente_id, "cliente_nome": cliente_nome,
        "ativo_id": ativo_id, "ativo_nome": ativo_nome,
        "periodo_inicio": periodo_inicio, "periodo_fim": periodo_fim,
        "modo": modo, "tipo_resumo": tipo_resumo,
        "indicadores": indicadores, "charts": charts,
        "pontos_gerencia": pontos, "acoes_prioritarias": acoes,
        "timeline": timeline, "report_ids_usados": report_ids_usados,
        "gerado_em": _dt.datetime.now(),
    }

    if salvar:
        from sheets import add_executive_summary
        summary_id = add_executive_summary(
            cliente_id=cliente_id, titulo=titulo, resumo_texto=texto,
            gerado_por_usuario_id=usuario_id, ativo_id=ativo_id,
            tipo_resumo=tipo_resumo, modo=modo,
            periodo_inicio=periodo_inicio.strftime("%d/%m/%Y"),
            periodo_fim=periodo_fim.strftime("%d/%m/%Y"),
            dados_usados=",".join(k for k, v in incluir.items() if v),
            report_ids_usados=",".join(report_ids_usados),
        )
        resultado["summary_id"] = summary_id

    return resultado


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORTAÇÃO — WORD (.docx)
# ═══════════════════════════════════════════════════════════════════════════════
# Reaproveita python-docx (já em requirements.txt — mesma lib de
# report_word_generator.py) para texto/tabelas editáveis, e matplotlib
# (também já em requirements.txt) só para renderizar os gráficos como
# imagens PNG embutidas — não há biblioteca de geração de PDF no projeto
# (pdfplumber/PyPDF2 só leem PDF); instalar uma nova para isso contrariaria
# o pedido explícito de evitar dependência pesada, então a exportação desta
# etapa continua Word + copiar texto, conforme o projeto já suporta.

_NAVY  = "0F1F3D"
_BLUE  = "2563EB"
_MUTED = "64748B"

_COR_HEX = {
    "Bom": "10B981", "Atenção": "F59E0B", "Crítico": "EF4444", "Urgente": "7C3AED",
    "Baixa": "10B981", "Moderada": "F59E0B", "Alta": "F97316", "Crítica": "EF4444",
    "Em dia": "10B981", "Próximas": "F59E0B", "Vencidas": "EF4444",
    "Concluídas": "2563EB", "Por condição": "8B5CF6",
    "Normal": "10B981",
}


def _mpl_bar_png(labels: list, values: list, cores: list, altura: float = 2.6) -> bytes | None:
    if not labels:
        return None
    import io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(5.6, altura), dpi=150)
    bars = ax.bar(labels, values, color=[f"#{c}" for c in cores])
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height(), str(v),
               ha="center", va="bottom", fontsize=9, color="#334155")
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.get_yaxis().set_visible(False)
    ax.tick_params(axis="x", labelsize=9, colors="#334155")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", transparent=True)
    plt.close(fig)
    return buf.getvalue()


def _mpl_pizza_png(labels: list, values: list, cores: list) -> bytes | None:
    if not labels:
        return None
    import io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(4.2, 3.4), dpi=150)
    ax.pie(values, labels=[f"{l} ({v})" for l, v in zip(labels, values)],
          colors=[f"#{c}" for c in cores], startangle=90,
          textprops={"fontsize": 9, "color": "#334155"})
    ax.set_aspect("equal")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", transparent=True)
    plt.close(fig)
    return buf.getvalue()


def _mpl_linha_png(pontos: list) -> bytes | None:
    if len(pontos) < 2:
        return None
    import io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    xs = [p[0] for p in pontos]
    ys = [p[1] for p in pontos]
    fig, ax = plt.subplots(figsize=(5.6, 2.4), dpi=150)
    ax.plot(xs, ys, marker="o", color="#2563EB", linewidth=2)
    ax.set_ylim(0, 100)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=8, colors="#334155")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", transparent=True)
    plt.close(fig)
    return buf.getvalue()


def gerar_resumo_executivo_word(resultado: dict) -> bytes:
    """Gera o .docx do resumo executivo — capa profissional, cards de
    indicadores, gráficos (imagens), ações prioritárias, principais pontos
    e conclusão. Texto e tabelas continuam editáveis; gráficos são imagens
    incorporadas (matplotlib, sem dependência nova)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx não instalado. Verifique requirements.txt.")
    import io, os

    def _rgb(h: str) -> "RGBColor":
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _heading(doc, texto, tamanho=13, cor=_BLUE):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(tamanho)
        r.font.color.rgb = _rgb(cor)
        return p

    ind      = resultado.get("indicadores", {})
    charts   = resultado.get("charts", {})
    pontos   = resultado.get("pontos_gerencia", [])
    acoes    = resultado.get("acoes_prioritarias", [])
    timeline = resultado.get("timeline", [])
    p_ini    = resultado["periodo_inicio"].strftime("%d/%m/%Y")
    p_fim    = resultado["periodo_fim"].strftime("%d/%m/%Y")
    gerado   = resultado.get("gerado_em")
    gerado_str = gerado.strftime("%d/%m/%Y %H:%M") if gerado else ""

    doc = Document()

    # ── Capa ──────────────────────────────────────────────────────────────
    logo_path = os.path.join(os.path.dirname(__file__), "logo.jpg")
    if os.path.exists(logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(1.3))
        except Exception:
            pass

    marca = doc.add_paragraph()
    marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_marca = marca.add_run("PRED.IO")
    r_marca.bold = True
    r_marca.font.size = Pt(11)
    r_marca.font.color.rgb = _rgb(_MUTED)

    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run("Resumo Executivo de Confiabilidade")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _rgb(_NAVY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        f"Cliente: {resultado.get('cliente_nome','')}"
        + (f"  ·  Ativo: {resultado.get('ativo_nome')}" if resultado.get("ativo_nome") else "")
        + f"\nPeríodo: {p_ini} a {p_fim}  ·  Gerado em {gerado_str}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = _rgb(_MUTED)

    intro_p = doc.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro_p.add_run(
        "Visão consolidada da condição dos ativos, prioridades de manutenção, "
        "relatórios técnicos, alertas e ações recomendadas no período."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.color.rgb = _rgb(_MUTED)

    doc.add_page_break()

    # ── Indicadores principais (tabela) ──────────────────────────────────────
    _heading(doc, "Indicadores Principais", 15, _NAVY)
    cards = [
        ("Ativos monitorados", ind.get("ativos_monitorados", 0)),
        ("Saúde média", f"{ind.get('saude_media')}/100" if ind.get("saude_media") is not None else "—"),
        ("Ativos em atenção", ind.get("ativos_atencao", 0)),
        ("Ativos críticos", ind.get("ativos_criticos", 0)),
        ("Manutenções vencidas", ind.get("manutencoes_vencidas", 0)),
        ("Itens GUT críticos", ind.get("itens_gut_criticos", 0)),
        ("Alertas críticos", ind.get("alertas_criticos", 0)),
        ("Chamados abertos", ind.get("chamados_abertos", 0)),
        ("Relatórios no período", ind.get("relatorios_publicados", 0)),
    ]
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i in range(0, len(cards), 2):
        row = tbl.add_row().cells
        row[0].text = cards[i][0]
        row[1].text = str(cards[i][1])
        if i + 1 < len(cards):
            row[2].text = cards[i + 1][0]
            row[3].text = str(cards[i + 1][1])
    doc.add_paragraph()

    # ── Gráficos ──────────────────────────────────────────────────────────
    _heading(doc, "Indicadores Gráficos", 15, _NAVY)

    saude = charts.get("saude_ativos", {})
    if saude:
        doc.add_paragraph("Saúde dos ativos").runs[0].bold = True
        png = _mpl_pizza_png(list(saude.keys()), list(saude.values()),
                             [_COR_HEX.get(l, "94A3B8") for l in saude.keys()])
        if png:
            doc.add_picture(io.BytesIO(png), width=Inches(3.2))

    gut_p = charts.get("gut_prioridade", {})
    if gut_p:
        ordem = ["Baixa", "Moderada", "Alta", "Crítica"]
        vals  = [gut_p.get(l, 0) for l in ordem]
        doc.add_paragraph("GUT por prioridade").runs[0].bold = True
        png = _mpl_bar_png(ordem, vals, [_COR_HEX[l] for l in ordem])
        if png:
            doc.add_picture(io.BytesIO(png), width=Inches(5.6))

    manut = charts.get("manutencoes_status", {})
    if manut:
        ordem = [l for l in ["Em dia", "Próximas", "Vencidas", "Concluídas", "Por condição"] if manut.get(l)]
        vals  = [manut[l] for l in ordem]
        doc.add_paragraph("Manutenções por status").runs[0].bold = True
        png = _mpl_bar_png(ordem, vals, [_COR_HEX[l] for l in ordem])
        if png:
            doc.add_picture(io.BytesIO(png), width=Inches(5.6))

    sev = charts.get("relatorios_severidade", {})
    if sev:
        ordem = [l for l in ["Normal", "Atenção", "Crítico", "Urgente"] if sev.get(l)]
        vals  = [sev[l] for l in ordem]
        doc.add_paragraph("Relatórios por severidade").runs[0].bold = True
        png = _mpl_bar_png(ordem, vals, [_COR_HEX[l] for l in ordem])
        if png:
            doc.add_picture(io.BytesIO(png), width=Inches(5.6))

    evolucao = charts.get("score_evolucao", [])
    png_evol = _mpl_linha_png(evolucao)
    if png_evol:
        doc.add_paragraph("Evolução do score de saúde").runs[0].bold = True
        doc.add_picture(io.BytesIO(png_evol), width=Inches(5.6))

    top5 = charts.get("top5_criticos", [])
    if top5:
        doc.add_paragraph("Top 5 ativos mais críticos").runs[0].bold = True
        t5 = doc.add_table(rows=1, cols=4)
        t5.style = "Light Grid Accent 1"
        hdr = t5.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Ativo", "Score", "GUT", "Prioridade"
        for it in top5:
            row = t5.add_row().cells
            row[0].text = str(it.get("ativo", ""))
            row[1].text = str(it.get("score", "—"))
            row[2].text = str(it.get("gut_score", "—"))
            row[3].text = str(it.get("gut_prioridade", "—"))

    doc.add_page_break()

    # ── Principais pontos para gerência ──────────────────────────────────────
    _heading(doc, "Principais Pontos para Gerência", 15, _NAVY)
    for p in pontos:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(p)

    doc.add_paragraph()

    # ── Ações prioritárias ────────────────────────────────────────────────────
    _heading(doc, "Ações Prioritárias", 15, _NAVY)
    if acoes:
        rank_label = {1: "GUT Crítica", 2: "GUT Alta", 3: "Manutenção Vencida",
                     4: "Alerta Crítico", 5: "Recomendação por Condição"}
        ta = doc.add_table(rows=1, cols=5)
        ta.style = "Light Grid Accent 1"
        hdr = ta.rows[0].cells
        for i, h in enumerate(["Prioridade", "Ação", "Ativo", "Prazo/Data", "Status"]):
            hdr[i].text = h
        for a in acoes:
            row = ta.add_row().cells
            row[0].text = rank_label.get(a["rank"], "")
            row[1].text = str(a.get("acao", ""))
            row[2].text = str(a.get("ativo_id", "") or "—")
            row[3].text = str(a.get("prazo", ""))[:16]
            row[4].text = str(a.get("status", "") or "—")
    else:
        doc.add_paragraph("Nenhuma ação prioritária identificada no período.")

    doc.add_paragraph()

    # ── Histórico resumido ────────────────────────────────────────────────────
    _heading(doc, "Histórico Resumido do Período", 15, _NAVY)
    if timeline:
        for data, tipo, titulo_ev in timeline:
            doc.add_paragraph(f"{data} — {tipo} — {titulo_ev}", style="List Bullet")
    else:
        doc.add_paragraph("Nenhum evento relevante registrado no período.")

    doc.add_page_break()

    # ── Conclusão (reaproveita a seção 10 do texto já montado) ────────────────
    _heading(doc, "Conclusão para Reunião", 15, _NAVY)
    texto = resultado.get("texto", "")
    if "10. CONCLUSÃO PARA REUNIÃO" in texto:
        bloco = texto.split("10. CONCLUSÃO PARA REUNIÃO", 1)[1]
        for linha in bloco.split("\n"):
            s = linha.strip()
            if not s or s == "Fonte: Pred.IO":
                continue
            if s.startswith("•"):
                doc.add_paragraph(s.lstrip("• ").strip(), style="List Bullet")
            else:
                doc.add_paragraph(s)

    rodape = doc.add_paragraph()
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rodape = rodape.add_run(TEXTO_OBRIGATORIO + "\n\nFonte: Pred.IO")
    r_rodape.italic = True
    r_rodape.font.size = Pt(9)
    r_rodape.font.color.rgb = _rgb(_MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
